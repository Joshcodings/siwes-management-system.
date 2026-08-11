import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_code_section(doc, title, filepath):
    doc.add_heading(title, level=2)
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
        except Exception as e:
            doc.add_paragraph(f"Error reading {filepath}: {str(e)}")
    else:
        doc.add_paragraph(f"File not found: {filepath}")

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Appendices', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Appendix A
    doc.add_heading('Appendix A: Frontend Application Source Code (React)', level=1)
    doc.add_paragraph("This appendix contains the core frontend implementation of the SIWES Management System, including the main application logic, layout, and styling.")
    
    add_code_section(doc, 'A.1 Main Application Component (src/App.tsx)', 'src/App.tsx')
    add_code_section(doc, 'A.2 Global Styles (src/index.css)', 'src/index.css')
    add_code_section(doc, 'A.3 Entry Point (src/main.tsx)', 'src/main.tsx')
    
    doc.add_page_break()
    
    # Appendix B
    doc.add_heading('Appendix B: Backend Server Source Code (Node.js)', level=1)
    doc.add_paragraph("This appendix contains the backend server implementation, handling API requests, database interactions, authentication, and file uploads.")
    
    add_code_section(doc, 'B.1 Express Server (server.ts)', 'server.ts')
    
    doc.add_page_break()
    
    # Appendix C
    doc.add_heading('Appendix C: AI Placement Engine Source Code (Python)', level=1)
    doc.add_paragraph("This appendix contains the AI-powered recommendation engine script, which scores students against available companies using natural language processing and geospatial data.")
    
    add_code_section(doc, 'C.1 AI Engine (ai_engine.py)', 'ai_engine.py')
    
    doc.add_page_break()
    
    # Appendix D
    doc.add_heading('Appendix D: Database Seeding Scripts', level=1)
    doc.add_paragraph("This appendix contains the scripts used to seed the database with initial mock data for testing and development.")
    
    add_code_section(doc, 'D.1 Database Seeder (seed_real_companies.ts)', 'seed_real_companies.ts')
    
    # Save the document
    output_path = r'C:\Users\hp\Downloads\SIWES_Management_System_Appendices.docx'
    doc.save(output_path)
    print(f"Successfully generated Appendices document at: {output_path}")

if __name__ == '__main__':
    main()
